"""Async document ingestion worker.

This worker handles the full document processing pipeline:

1. Fetch the raw document (from S3 or local storage)
2. Extract text content (PDF, DOCX, HTML, plain text)
3. Split the content into chunks
4. Generate embeddings for each chunk
5. Persist chunks with embeddings to the database

The worker processes documents whose status is ``PENDING`` and
transitions them through ``PROCESSING`` -> ``INDEXED`` (or ``FAILED``).
"""

from __future__ import annotations

import asyncio
import hashlib
import io
import uuid

import structlog
from sqlalchemy import select, update
from sqlalchemy.ext.asyncio import AsyncSession

from autonomocx.ai.rag.embeddings import EmbeddingService
from autonomocx.models.knowledge import Document, DocumentChunk, DocumentStatus, KnowledgeBase
from autonomocx.services.storage_service import storage_service

logger = structlog.get_logger(__name__)


class DocumentProcessor:
    """Processes uploaded documents into searchable text chunks with embeddings."""

    def __init__(self, chunk_size: int = 512, chunk_overlap: int = 64) -> None:
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self._embedder: EmbeddingService | None = None

    def _get_embedder(self) -> EmbeddingService:
        if self._embedder is None:
            self._embedder = EmbeddingService()
        return self._embedder

    async def process_document(self, db: AsyncSession, document_id: uuid.UUID) -> None:
        """Run the full ingestion pipeline for a single document.

        Flow:
            1. Mark document as PROCESSING
            2. Extract text from the source file
            3. Split text into overlapping chunks
            4. Generate embeddings for chunks
            5. Persist chunks with embeddings to the database
            6. Update document status to INDEXED
            7. Update knowledge base counters
        """
        # 1. Load and mark as processing
        result = await db.execute(select(Document).where(Document.id == document_id))
        document = result.scalar_one_or_none()
        if document is None:
            logger.error("document_not_found", document_id=str(document_id))
            return

        document.status = DocumentStatus.PROCESSING
        await db.flush()
        logger.info("document_processing_started", document_id=str(document_id))

        try:
            # 2. Extract text content
            raw_text = await self._extract_text(document)
            if not raw_text.strip():
                raise ValueError("No text content could be extracted from the document.")

            # 3. Compute content hash for deduplication
            content_hash = hashlib.sha256(raw_text.encode()).hexdigest()
            document.content_hash = content_hash

            # 4. Split into chunks
            chunks = self._split_into_chunks(raw_text)
            logger.info(
                "document_chunked",
                document_id=str(document_id),
                num_chunks=len(chunks),
            )

            # 5. Generate embeddings
            embedder = self._get_embedder()
            embeddings = await embedder.embed_batch(chunks)
            logger.info(
                "document_embedded",
                document_id=str(document_id),
                num_embeddings=len(embeddings),
            )

            # 6. Persist chunks with embeddings
            for idx, (chunk_text, embedding) in enumerate(zip(chunks, embeddings)):
                chunk = DocumentChunk(
                    id=uuid.uuid4(),
                    document_id=document.id,
                    knowledge_base_id=document.knowledge_base_id,
                    chunk_index=idx,
                    content=chunk_text,
                    embedding=embedding,
                    token_count=self._estimate_token_count(chunk_text),
                    metadata_={"chunk_index": idx, "total_chunks": len(chunks)},
                )
                db.add(chunk)

            # 7. Update document status
            document.status = DocumentStatus.INDEXED
            document.chunk_count = len(chunks)
            await db.flush()

            # 8. Update knowledge base counters
            await db.execute(
                update(KnowledgeBase)
                .where(KnowledgeBase.id == document.knowledge_base_id)
                .values(
                    total_chunks=KnowledgeBase.total_chunks + len(chunks),
                )
            )

            await db.commit()
            logger.info(
                "document_processing_complete",
                document_id=str(document_id),
                chunks_created=len(chunks),
            )

        except Exception as exc:
            await db.rollback()
            # Re-fetch document after rollback
            result = await db.execute(select(Document).where(Document.id == document_id))
            document = result.scalar_one_or_none()
            if document is not None:
                document.status = DocumentStatus.FAILED
                document.error_message = str(exc)[:1000]
                await db.commit()
            logger.exception(
                "document_processing_failed",
                document_id=str(document_id),
                error=str(exc),
            )

    async def _extract_text(self, document: Document) -> str:
        """Extract text content from a document based on its file type."""
        file_type = (document.file_type or "").lower()

        if file_type == "pdf":
            return await self._extract_pdf(document)
        elif file_type in ("docx", "doc"):
            return await self._extract_docx(document)
        elif file_type in ("html", "htm"):
            return await self._extract_html(document)
        elif file_type in ("txt", "md", "csv", "json"):
            return await self._extract_plain_text(document)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    async def _extract_pdf(self, document: Document) -> str:
        """Extract text from a PDF file using pypdf."""
        content = await self._download_file(document)
        logger.info("extracting_pdf", file_path=document.file_path)

        def _read_pdf(data: bytes) -> str:
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(data))
            pages: list[str] = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)
            return "\n\n".join(pages)

        return await asyncio.to_thread(_read_pdf, content)

    async def _extract_docx(self, document: Document) -> str:
        """Extract text from a DOCX file using python-docx."""
        content = await self._download_file(document)
        logger.info("extracting_docx", file_path=document.file_path)

        def _read_docx(data: bytes) -> str:
            from docx import Document as DocxDocument

            doc = DocxDocument(io.BytesIO(data))
            parts: list[str] = []
            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            return "\n\n".join(parts)

        return await asyncio.to_thread(_read_docx, content)

    async def _extract_html(self, document: Document) -> str:
        """Extract text from HTML using BeautifulSoup."""
        content = await self._download_file(document)
        logger.info("extracting_html", source_url=document.source_url)

        def _read_html(data: bytes) -> str:
            from bs4 import BeautifulSoup

            soup = BeautifulSoup(data, "html.parser")
            # Remove script and style elements
            for tag in soup(["script", "style", "nav", "footer", "header"]):
                tag.decompose()
            text = soup.get_text(separator="\n", strip=True)
            # Collapse multiple newlines
            lines = [line.strip() for line in text.splitlines() if line.strip()]
            return "\n\n".join(lines)

        return await asyncio.to_thread(_read_html, content)

    async def _extract_plain_text(self, document: Document) -> str:
        """Read plain text content directly."""
        content = await self._download_file(document)
        logger.info("extracting_text", file_path=document.file_path)
        return content.decode("utf-8", errors="replace")

    async def _download_file(self, document: Document) -> bytes:
        """Download file content from storage."""
        if not document.file_path:
            raise ValueError(f"Document {document.id} has no file_path set.")
        return await storage_service.download_file(document.file_path)

    def _split_into_chunks(self, text: str) -> list[str]:
        """Split text into overlapping chunks of approximately ``chunk_size`` tokens.

        Uses a simple word-boundary-aware splitting strategy.
        """
        words = text.split()
        if not words:
            return []

        # Rough approximation: 1 token ~ 0.75 words
        words_per_chunk = max(1, int(self.chunk_size * 0.75))
        overlap_words = max(0, int(self.chunk_overlap * 0.75))

        chunks: list[str] = []
        start = 0
        while start < len(words):
            end = start + words_per_chunk
            chunk = " ".join(words[start:end])
            chunks.append(chunk)
            start += words_per_chunk - overlap_words

        return chunks

    @staticmethod
    def _estimate_token_count(text: str) -> int:
        """Estimate token count using a rough heuristic (4 chars per token)."""
        return max(1, len(text) // 4)

    async def process_pending_documents(self, db: AsyncSession) -> int:
        """Find and process all documents in ``PENDING`` status.

        Returns the number of documents processed.
        """
        result = await db.execute(select(Document).where(Document.status == DocumentStatus.PENDING))
        documents = result.scalars().all()

        processed = 0
        for doc in documents:
            await self.process_document(db, doc.id)
            processed += 1

        logger.info("pending_documents_processed", count=processed)
        return processed
